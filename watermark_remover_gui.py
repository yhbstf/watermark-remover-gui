# -*- coding: utf-8 -*-
"""
Watermark Remover GUI v5 — multi-tool selection, zoom/pan,
Word/PDF/PPT/Excel support, batch modes, templates, CLI.
"""

import argparse
import json
import locale
import os
import shutil
import subprocess
import sys
import time
import zipfile
from pathlib import Path

import cv2
import numpy as np
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from PIL import Image, ImageTk
from docx import Document


# Optional dependencies — features gracefully degrade if unavailable.
try:
    import fitz  # PyMuPDF
    HAVE_PDF = True
except Exception:
    HAVE_PDF = False

try:
    from tkinterdnd2 import TkinterDnD, DND_FILES
    HAVE_DND = True
except Exception:
    HAVE_DND = False

try:
    import pytesseract
    HAVE_OCR = True
except Exception:
    HAVE_OCR = False


TEMP_ROOT = Path("D:/temp/watermark_remover_temp") if sys.platform == "win32" \
    else Path.home() / ".cache" / "watermark-remover" / "tmp"
BATCH_TEMP_ROOT = Path("D:/temp/watermark_remover_batch") if sys.platform == "win32" \
    else Path.home() / ".cache" / "watermark-remover" / "batch"
EXTRACT_ROOT = Path("D:/temp/docx_temp_extract") if sys.platform == "win32" \
    else Path.home() / ".cache" / "watermark-remover" / "extract"
BATCH_EXTRACT_ROOT = Path("D:/temp/docx_temp_extract_batch") if sys.platform == "win32" \
    else Path.home() / ".cache" / "watermark-remover" / "extract_batch"

CONFIG_PATH = Path.home() / ".watermark-remover.json"
TEMPLATE_EXT = ".wmrtpl"

MODELS = ["lama", "mat", "fcf", "ldm", "manga"]
IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tif", ".tiff")
OOXML_IMAGE_PREFIXES = ("word/media/", "ppt/media/", "xl/media/")


# ----------------------------------------------------------- config
def load_config():
    if CONFIG_PATH.exists():
        try:
            return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}


def save_config(cfg):
    try:
        CONFIG_PATH.write_text(
            json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass


CONFIG = load_config()


# ----------------------------------------------------------- i18n
def _detect_lang():
    env = os.environ.get("WMR_LANG", "").lower()
    if env in ("en", "zh"):
        return env
    if CONFIG.get("language") in ("en", "zh"):
        return CONFIG["language"]
    for var in ("LC_ALL", "LC_MESSAGES", "LANG"):
        val = os.environ.get(var, "")
        if val.lower().startswith("zh"):
            return "zh"
    try:
        current = (locale.getlocale()[0] or "").lower()
        if current.startswith("zh") or "chinese" in current:
            return "zh"
    except Exception:
        pass
    if sys.platform == "win32":
        try:
            import ctypes
            lang_id = ctypes.windll.kernel32.GetUserDefaultUILanguage()
            if (lang_id & 0x3FF) == 0x04:
                return "zh"
        except Exception:
            pass
    return "en"


LANG = _detect_lang()

_ZH = {
    "Watermark Remover v5": "水印去除工具 v5",
    "Select Image": "选择图片",

    "File": "文件",
    "Edit": "编辑",
    "View": "视图",
    "Tools": "工具",
    "Template": "模板",
    "Open Image File": "打开图片文件",
    "Open from Word": "从 Word 打开",
    "Open from PDF": "从 PDF 打开",
    "Open from PPT": "从 PPT 打开",
    "Open from Excel": "从 Excel 打开",
    "Recent Files": "最近打开",
    "(no recent files)": "（无最近文件）",
    "Clear Recent": "清空最近列表",
    "Batch Folder...": "批处理文件夹...",
    "Exit": "退出",

    "Undo (Ctrl+Z)": "撤销 (Ctrl+Z)",
    "Redo (Ctrl+Y)": "重做 (Ctrl+Y)",

    "Save Template...": "保存模板...",
    "Load Template...": "加载模板...",

    "Compare (before/after)": "前后对比",
    "Auto-detect watermarks": "自动检测水印",
    "Detect text (OCR)": "检测文字水印 (OCR)",
    "Select by color...": "按颜色选择...",

    "Preview - drag to box, Shift+drag/middle-drag to pan, wheel to zoom":
        "预览 — 拖拽画框，Shift+拖 / 中键拖 平移，滚轮缩放",
    "Control Panel": "控制面板",
    "Current file:": "当前文件：",
    "Regions:": "选区：",
    "Status:": "状态：",
    "Images:": "图片：",
    "Preview:": "预览：",
    "Select image to edit:": "选择要编辑的图片：",

    "Tool:": "工具：",
    "Rect": "矩形",
    "Brush": "画笔",
    "Polygon": "多边形",
    "Color": "颜色",
    "Brush size:": "画笔大小：",
    "Feather:": "羽化：",
    "Model:": "模型：",
    "Color tolerance:": "颜色容差：",

    "Undo Last Region": "撤销最后一个选区",
    "Clear All Regions (Esc)": "清空所有选区 (Esc)",
    "Remove Watermark": "去除水印",
    "Batch: Apply to ALL Word Images": "批量：应用到所有 Word 图片",
    "Reset to Original (Ctrl+Z)": "恢复原图",
    "Zoom -": "缩小",
    "Fit (0)": "适应 (0)",
    "Zoom +": "放大",
    "Save to Word": "保存到 Word",
    "Export PNG": "导出 PNG",
    "Confirm": "确定",
    "Cancel": "取消",
    "OK": "确定",

    "Error": "错误",
    "Warning": "警告",
    "Success": "成功",
    "Info": "提示",
    "Batch Remove": "批量去除",

    "Ready.\n"
    "1. File -> Open image / Open from Word\n"
    "2. Draw selection (Rect/Brush/Polygon/Color)\n"
    "3. Remove Watermark  OR  Batch across Word\n"
    "4. Save to Word / Export PNG\n"
    "Wheel = zoom at cursor\n"
    "Shift+drag or middle-drag = pan\n"
    "+/-/0 = zoom in/out/fit, Esc = clear, Ctrl+Z/Y = undo/redo":
        "就绪。\n"
        "1. 文件 → 打开图片 / 从 Word 打开\n"
        "2. 画选区（矩形/画笔/多边形/颜色）\n"
        "3. 去除水印 或 批量处理 Word\n"
        "4. 保存到 Word / 导出 PNG\n"
        "滚轮：以鼠标为中心缩放\n"
        "Shift+拖 或 中键拖：平移\n"
        "+/-/0：放大/缩小/适应，Esc：清空，Ctrl+Z/Y：撤销/重做",

    "Opened: ": "已打开：",
    "Size: {}x{}": "尺寸：{}x{}",
    "Read error: ": "读取错误：",
    "Cannot open image:\n": "无法打开图片：\n",
    "Cannot open Word file: ": "无法打开 Word 文件：",
    "Cannot open PDF: ": "无法打开 PDF：",
    "Cannot open archive: ": "无法打开文件：",
    "No images found": "未找到图片",
    "Found {} images": "找到 {} 张图片",
    "Please select an image": "请选择一张图片",
    "Cannot decode image": "无法解码图片",
    "Loaded: ": "已加载：",
    "Please open an image first": "请先打开图片",
    "Draw at least one region first": "请先画至少一个选区",
    "Selection is empty": "选区为空",
    "Removing watermark...": "正在去除水印...",
    "Done.": "完成。",
    "Watermark removed": "水印已去除",
    "iopaint failed:\n": "iopaint 失败：\n",
    "iopaint timed out (>10min)": "iopaint 超时（>10 分钟）",
    "Failed: ": "失败：",
    "No image to save": "没有可保存的图片",
    "This image was not opened from a document. Use Export PNG instead.":
        "此图片未从文档打开，请使用“导出 PNG”。",
    "Saved to {} (backup: {})": "已保存到 {}（备份：{}）",
    "Saved to:\n": "已保存到：\n",
    "Failed to save: ": "保存失败：",
    "Exported: ": "已导出：",
    "Failed to encode image": "图片编码失败",
    "Removed last region": "已撤销最后一个选区",
    "Cleared all regions": "已清空所有选区",
    "Reset to original": "已恢复原图",
    "Undone": "已撤销",
    "Redone": "已重做",
    "Nothing to undo": "无可撤销",
    "Nothing to redo": "无可重做",

    "Type: Image file\n": "类型：图片文件\n",
    "Type: {} document\n{}\nrId: {}": "类型：{} 文档\n{}\nrId：{}",

    "Open a Word document first (File -> Open from Word)":
        "请先打开文档（文件 → 从 ... 打开）",
    "Draw at least one region on the current image first":
        "请先在当前图片上画至少一个选区",
    "Apply the selection you drew to ALL {} images in:\n"
    "{}\n\n"
    "Mask scales proportionally per image.\n"
    "A .backup copy of the file will be created.\n\nContinue?":
        "将你画的选区应用到以下文件中所有 {} 张图片：\n"
        "{}\n\n"
        "选区会按每张图片的尺寸等比缩放。\n"
        "会创建 .backup 备份。\n\n是否继续？",
    "Backup: ": "备份：",
    "No images were processed.": "没有图片被处理。",
    "Batch complete.\nProcessed: {}\nFailed: {}\nBackup: {}":
        "批量完成。\n已处理：{}\n失败：{}\n备份：{}",
    "\n\nFailed images:\n": "\n\n失败的图片：\n",
    "\n  ... and {} more": "\n  ... 还有 {} 张",

    "Total regions: {}\n": "选区总数：{}\n",
    "drawing: x=[{}:{}] y=[{}:{}]\n": "绘制中：x=[{}:{}] y=[{}:{}]\n",

    "  skip (cannot decode)": "  跳过（无法解码）",
    "  skip (empty mask)": "  跳过（选区为空）",
    "  encode failed": "  编码失败",
    "  ok": "  成功",
    "  FAILED: iopaint: ": "  失败：iopaint：",
    "  FAILED: ": "  失败：",
    "[{}/{}] {}": "[{}/{}] {}",

    "Language switched to: ": "已切换语言：",

    "PyMuPDF not installed. Run: pip install pymupdf":
        "未安装 PyMuPDF。请运行：pip install pymupdf",
    "pytesseract / Tesseract not installed. Run: pip install pytesseract (and install Tesseract)":
        "未安装 pytesseract / Tesseract。请运行：pip install pytesseract（并安装 Tesseract）",
    "Drag-and-drop disabled (tkinterdnd2 not installed)":
        "拖放已禁用（未安装 tkinterdnd2）",

    "Detected {} candidate region(s)": "检测到 {} 个候选区域",
    "No watermarks detected": "未检测到水印",
    "Click a pixel to sample color": "点击像素以取样颜色",
    "Color tolerance": "颜色容差",

    "Batch Folder": "批处理文件夹",
    "Select folder": "选择文件夹",
    "Process {} image(s) in {}?": "处理 {} 中的 {} 张图片？",
    "Processed {} image(s). Output folder: {}": "已处理 {} 张图片。输出目录：{}",

    "Save template": "保存模板",
    "Load template": "加载模板",
    "Template saved": "模板已保存",
    "Template loaded": "模板已加载",
    "Cannot load template: ": "无法加载模板：",

    "Before": "修改前",
    "After": "修改后",
    "Compare": "对比",

    "Progress: {}/{}": "进度：{}/{}",
    "Processing...": "处理中...",
    "Drag and drop files here": "拖入文件到此",

    "Open Video...": "打开视频...",
    "Process Video": "处理视频",
    "Video": "视频",
    "Method: (fast) cv2 inpaint / (slow) LaMa": "方法：(快) cv2 补全 / (慢) LaMa",
    "Export video to:": "导出视频到：",
    "Processing frame {}/{}": "处理帧 {}/{}",
    "Video done. Output: {}": "视频完成。输出：{}",
    "Cannot open video": "无法打开视频",
    "Quick Templates": "快速模板",
    "Template applied": "模板已应用",
    "Preview result": "预览结果",
    "Apply": "应用",
    "Deep Clean": "深度清理",
    "Show Loupe (L)": "显示放大镜 (L)",
    "Loupe": "放大镜",
    "Hotkeys: R rect / B brush / P polygon / C color / L loupe / Del delete region":
        "快捷键: R 矩形 / B 画笔 / P 多边形 / C 颜色 / L 放大镜 / Del 删除选区",
    "Nothing detected after pass {}. Stopping.": "第 {} 轮后未检测到残留，停止。",
    "Deep clean pass {}/{}": "深度清理：第 {}/{} 轮",
    "Selected region #{}": "已选中选区 #{}",
    "Deleted region": "已删除选区",
}


def t(s):
    if LANG == "zh":
        return _ZH.get(s, s)
    return s


# Each template is a list of (x1, y1, x2, y2) in relative 0..1 coords
# of the image. Positions are rough heuristics; users can adjust.
PLATFORM_TEMPLATES = {
    "豆包AI (右下)": [(0.72, 0.92, 0.99, 0.99)],
    "抖音 / 用户ID (右下)": [(0.60, 0.90, 0.99, 0.99)],
    "抖音 / 大 LOGO (左上)": [(0.01, 0.01, 0.18, 0.10)],
    "快手 (右下用户名)": [(0.58, 0.91, 0.99, 0.99)],
    "小红书 (右下)": [(0.65, 0.92, 0.99, 0.99)],
    "B 站 / bilibili (右上)": [(0.72, 0.01, 0.99, 0.10)],
    "微博 (底部条)": [(0.00, 0.94, 1.00, 0.99)],
    "即梦 AI (右下)": [(0.68, 0.92, 0.99, 0.99)],
    "Midjourney 水印 (左下)": [(0.01, 0.92, 0.25, 0.99)],
    "DALL·E 彩色条 (右下)": [(0.80, 0.92, 0.99, 0.99)],
    "Getty Images (居中大水印)": [(0.05, 0.35, 0.95, 0.65)],
    "Shutterstock (斜向)": [(0.00, 0.20, 1.00, 0.80)],
}

VIDEO_EXTS = (".mp4", ".avi", ".mov", ".mkv", ".webm", ".m4v", ".flv")


# --------------------------------------------- OOXML archive helpers
def is_ooxml_image_part(partname):
    return partname.lstrip("/").startswith(OOXML_IMAGE_PREFIXES)


def list_ooxml_images(path):
    """Return [(partname, bytes)] for image parts inside an OOXML file."""
    out = []
    with zipfile.ZipFile(path, "r") as zf:
        for info in zf.infolist():
            if is_ooxml_image_part(info.filename):
                with zf.open(info) as f:
                    out.append((info.filename, f.read()))
    return out


def replace_ooxml_images(path, replacements, extract_dir):
    """Replace image parts in an OOXML zip; `replacements` is {partname: bytes}."""
    if extract_dir.exists():
        shutil.rmtree(extract_dir)
    extract_dir.mkdir(parents=True)
    with zipfile.ZipFile(path, "r") as zf:
        zf.extractall(extract_dir)
    for partname, data in replacements.items():
        p = extract_dir / partname.lstrip("/")
        p.parent.mkdir(parents=True, exist_ok=True)
        with open(p, "wb") as f:
            f.write(data)
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for root_dir, _, files in os.walk(extract_dir):
            for file in files:
                fp = Path(root_dir) / file
                zf.write(fp, fp.relative_to(extract_dir))
    shutil.rmtree(extract_dir)


# --------------------------------------------- auto-detect (MSER based)
def _merge_overlapping_boxes(boxes, dx=15, dy=10):
    """Merge boxes that are close or overlap after expanding by (dx, dy)."""
    if not boxes:
        return boxes
    bs = [list(b) for b in boxes]
    changed = True
    while changed:
        changed = False
        i = 0
        while i < len(bs):
            j = i + 1
            while j < len(bs):
                a, b = bs[i], bs[j]
                ax1, ay1, ax2, ay2 = a[0] - dx, a[1] - dy, a[2] + dx, a[3] + dy
                if not (b[2] < ax1 or b[0] > ax2 or b[3] < ay1 or b[1] > ay2):
                    a[0] = min(a[0], b[0])
                    a[1] = min(a[1], b[1])
                    a[2] = max(a[2], b[2])
                    a[3] = max(a[3], b[3])
                    bs.pop(j)
                    changed = True
                else:
                    j += 1
            i += 1
    return bs


def auto_detect_regions(cv_img):
    """
    MSER-based candidate detector for watermarks.
    Down-scales for speed, runs MSER on both polarities (dark- and
    light-on-background text), merges stroke-level boxes into word-level
    ones, then scores by corner proximity and size. Returns top 3.
    """
    if cv_img is None:
        return []
    orig_h, orig_w = cv_img.shape[:2]
    max_side = 1200
    s = min(1.0, max_side / max(orig_h, orig_w))
    small = cv2.resize(cv_img, (int(orig_w * s), int(orig_h * s))) if s < 1.0 else cv_img
    h, w = small.shape[:2]
    gray = cv2.cvtColor(small, cv2.COLOR_BGR2GRAY)

    try:
        mser = cv2.MSER_create()
        for attr, val in (("setDelta", 5),
                          ("setMinArea", 40),
                          ("setMaxArea", max(200, int(h * w * 0.01)))):
            try:
                getattr(mser, attr)(val)
            except Exception:
                pass
        regions_a, _ = mser.detectRegions(gray)
        regions_b, _ = mser.detectRegions(255 - gray)
        all_regions = list(regions_a) + list(regions_b)
    except Exception:
        all_regions = []

    raw = []
    for region in all_regions:
        x, y, rw, rh = cv2.boundingRect(region.reshape(-1, 1, 2))
        if rw < 6 or rh < 6:
            continue
        if rw > w * 0.6 or rh > h * 0.6:
            continue
        raw.append([x, y, x + rw, y + rh])

    merged = _merge_overlapping_boxes(
        raw, dx=max(6, int(w * 0.02)), dy=max(4, int(h * 0.015)))

    scored = []
    for (x1, y1, x2, y2) in merged:
        bw = x2 - x1
        bh = y2 - y1
        area = bw * bh
        if area < 400 or area > w * h * 0.25:
            continue
        if bw < 20 and bh < 20:
            continue
        ratio = bw / max(1, bh)
        if ratio > 40 or ratio < 0.04:
            continue
        cx = (x1 + x2) / 2.0
        cy = (y1 + y2) / 2.0
        # Distance to nearest image corner (watermarks usually sit near edges).
        d_corner = min(
            (cx ** 2 + cy ** 2) ** 0.5,
            ((w - cx) ** 2 + cy ** 2) ** 0.5,
            (cx ** 2 + (h - cy) ** 2) ** 0.5,
            ((w - cx) ** 2 + (h - cy) ** 2) ** 0.5,
        )
        corner_score = 1.0 - min(1.0, d_corner / (min(h, w) / 2.0 + 1e-6))
        size_score = min(1.0, area / (w * h * 0.03))
        score = corner_score * 2.5 + size_score * 0.5
        scored.append((score, x1, y1, x2, y2))

    scored.sort(reverse=True)
    inv = 1.0 / s if s < 1.0 else 1.0
    out = []
    for _, x1, y1, x2, y2 in scored[:3]:
        # Give the box a small margin so feathering/inpaint has room.
        pad_x = max(2, int((x2 - x1) * 0.05))
        pad_y = max(2, int((y2 - y1) * 0.05))
        out.append((
            int(max(0, (x1 - pad_x) * inv)),
            int(max(0, (y1 - pad_y) * inv)),
            int(min(orig_w, (x2 + pad_x) * inv)),
            int(min(orig_h, (y2 + pad_y) * inv)),
        ))
    return out


# --------------------------------------------- main class
class WatermarkRemover:
    SHIFT_MASK = 0x0001
    CTRL_MASK = 0x0004

    def __init__(self, root):
        self.root = root
        self.root.title(t("Watermark Remover v5"))
        geometry = CONFIG.get("geometry", "1400x900")
        self.root.geometry(geometry)

        self.cv_image = None
        self.original_image = None
        self.photo = None

        # Selection: rects stay as a separate list for clarity; an extra
        # pixel mask carries brush/polygon/color/auto-detect additions.
        self.rects = []
        self.in_progress = None
        self.drawing = False
        self.extra_mask = None  # HxW uint8 aligned with cv_image
        self.polygon_points = []

        # View.
        self.base_scale = 1.0
        self.zoom = 1.0
        self.pan_x = 0.0
        self.pan_y = 0.0
        self.panning = False
        self.pan_anchor = None

        # Document source.
        self.source_kind = None  # "image" | "word" | "pdf" | "pptx" | "xlsx"
        self.source_path = None
        self.source_extra = None  # rId for ooxml, xref for pdf, page num etc.
        self.word_doc = None
        self.pdf_doc = None
        self._file_desc = None

        # Tool mode and settings.
        self.tool_mode = tk.StringVar(value="rect")
        self.mask_op = tk.StringVar(value="add")  # add | sub
        self.brush_size = tk.IntVar(value=CONFIG.get("brush_size", 20))
        self.feather = tk.IntVar(value=CONFIG.get("feather", 0))
        self.color_tol = tk.IntVar(value=CONFIG.get("color_tol", 24))
        self.model = tk.StringVar(value=CONFIG.get("model", "lama"))
        self._color_pick_pending = False

        # Undo/redo.
        self.undo_stack = []
        self.redo_stack = []

        # Rect editing: currently-selected rect index + drag state.
        self.selected_rect_idx = None
        self.rect_drag = None  # (orig_rect, start_ix, start_iy)

        # Loupe window.
        self.loupe_win = None
        self.loupe_canvas = None
        self.loupe_photo = None

        # Video source state.
        self.video_path = None

        self._build_menu()
        self._build_layout()
        self._bind_events()
        self._setup_dnd()
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)

        self.log(t(
            "Ready.\n"
            "1. File -> Open image / Open from Word\n"
            "2. Draw selection (Rect/Brush/Polygon/Color)\n"
            "3. Remove Watermark  OR  Batch across Word\n"
            "4. Save to Word / Export PNG\n"
            "Wheel = zoom at cursor\n"
            "Shift+drag or middle-drag = pan\n"
            "+/-/0 = zoom in/out/fit, Esc = clear, Ctrl+Z/Y = undo/redo"
        ))
        if not HAVE_DND:
            self.log(t("Drag-and-drop disabled (tkinterdnd2 not installed)"))

    # -------------------- UI scaffolding
    def _build_menu(self):
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label=t("File"), menu=file_menu)
        file_menu.add_command(label=t("Open Image File"), command=self.open_image)
        file_menu.add_command(label=t("Open from Word"),
                              command=lambda: self.open_from_ooxml("word"))
        file_menu.add_command(label=t("Open from PPT"),
                              command=lambda: self.open_from_ooxml("pptx"))
        file_menu.add_command(label=t("Open from Excel"),
                              command=lambda: self.open_from_ooxml("xlsx"))
        file_menu.add_command(label=t("Open from PDF"),
                              command=self.open_from_pdf)
        file_menu.add_separator()
        file_menu.add_command(label=t("Batch Folder..."), command=self.batch_folder)
        file_menu.add_separator()
        self.recent_menu = tk.Menu(file_menu, tearoff=0)
        file_menu.add_cascade(label=t("Recent Files"), menu=self.recent_menu)
        self._rebuild_recent_menu()
        file_menu.add_separator()
        file_menu.add_command(label=t("Exit"), command=self._on_close)

        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label=t("Edit"), menu=edit_menu)
        edit_menu.add_command(label=t("Undo (Ctrl+Z)"), command=self.undo)
        edit_menu.add_command(label=t("Redo (Ctrl+Y)"), command=self.redo)

        tools_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label=t("Tools"), menu=tools_menu)
        tools_menu.add_command(label=t("Auto-detect watermarks"),
                               command=self.auto_detect)
        tools_menu.add_command(label=t("Detect text (OCR)"), command=self.ocr_detect)
        tools_menu.add_command(label=t("Select by color..."), command=self.begin_color_pick)
        tools_menu.add_separator()
        tools_menu.add_command(label=t("Deep Clean"), command=self.deep_clean)
        tools_menu.add_command(label=t("Show Loupe (L)"), command=self.toggle_loupe)
        tools_menu.add_command(label=t("Compare (before/after)"),
                               command=self.show_compare)

        video_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label=t("Video"), menu=video_menu)
        video_menu.add_command(label=t("Open Video..."), command=self.open_video)
        video_menu.add_command(label=t("Process Video"), command=self.process_video)

        tpl_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label=t("Template"), menu=tpl_menu)
        tpl_menu.add_command(label=t("Save Template..."), command=self.save_template)
        tpl_menu.add_command(label=t("Load Template..."), command=self.load_template)
        tpl_menu.add_separator()
        quick_menu = tk.Menu(tpl_menu, tearoff=0)
        tpl_menu.add_cascade(label=t("Quick Templates"), menu=quick_menu)
        for name in PLATFORM_TEMPLATES:
            quick_menu.add_command(
                label=name,
                command=lambda n=name: self.apply_platform_template(n))

        lang_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Language / 语言", menu=lang_menu)
        lang_menu.add_command(label="English", command=lambda: self.set_language("en"))
        lang_menu.add_command(label="中文", command=lambda: self.set_language("zh"))

    def _rebuild_recent_menu(self):
        self.recent_menu.delete(0, tk.END)
        recents = CONFIG.get("recent", [])
        if not recents:
            self.recent_menu.add_command(label=t("(no recent files)"), state="disabled")
        else:
            for p in recents[:10]:
                name = Path(p).name
                self.recent_menu.add_command(
                    label=name, command=lambda x=p: self.open_recent(x))
            self.recent_menu.add_separator()
            self.recent_menu.add_command(label=t("Clear Recent"),
                                         command=self.clear_recent)

    def _build_layout(self):
        main_frame = tk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        left = tk.Frame(main_frame)
        left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tk.Label(
            left,
            text=t("Preview - drag to box, Shift+drag/middle-drag to pan, wheel to zoom"),
            font=("Arial", 11, "bold"),
        ).pack()
        self.canvas = tk.Canvas(left, bg="gray", cursor="tcross")
        self.canvas.pack(fill=tk.BOTH, expand=True)

        right = tk.Frame(main_frame, width=340)
        right.pack(side=tk.RIGHT, fill=tk.BOTH, padx=10)
        tk.Label(right, text=t("Control Panel"), font=("Arial", 12, "bold")).pack()

        tk.Label(right, text=t("Current file:"), font=("Arial", 10)).pack(
            anchor="w", pady=(10, 5))
        self.file_text = tk.Text(right, height=2, width=44)
        self.file_text.pack(anchor="w", padx=5)

        # --- Tool selector
        tool_frame = tk.LabelFrame(right, text=t("Tool:"), padx=5, pady=5)
        tool_frame.pack(fill=tk.X, pady=(10, 5))
        for val, label in (("rect", t("Rect")), ("brush", t("Brush")),
                           ("polygon", t("Polygon")), ("color", t("Color"))):
            tk.Radiobutton(tool_frame, text=label, variable=self.tool_mode,
                           value=val).pack(side=tk.LEFT, padx=2)

        # Brush size
        bs_frame = tk.Frame(right)
        bs_frame.pack(fill=tk.X, pady=2)
        tk.Label(bs_frame, text=t("Brush size:"), width=10, anchor="w").pack(side=tk.LEFT)
        tk.Scale(bs_frame, from_=2, to=80, orient=tk.HORIZONTAL,
                 variable=self.brush_size).pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Feather
        ft_frame = tk.Frame(right)
        ft_frame.pack(fill=tk.X, pady=2)
        tk.Label(ft_frame, text=t("Feather:"), width=10, anchor="w").pack(side=tk.LEFT)
        tk.Scale(ft_frame, from_=0, to=30, orient=tk.HORIZONTAL,
                 variable=self.feather).pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Color tolerance
        ct_frame = tk.Frame(right)
        ct_frame.pack(fill=tk.X, pady=2)
        tk.Label(ct_frame, text=t("Color tolerance:"), width=10, anchor="w").pack(side=tk.LEFT)
        tk.Scale(ct_frame, from_=0, to=80, orient=tk.HORIZONTAL,
                 variable=self.color_tol).pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Model
        mdl_frame = tk.Frame(right)
        mdl_frame.pack(fill=tk.X, pady=(4, 8))
        tk.Label(mdl_frame, text=t("Model:"), width=10, anchor="w").pack(side=tk.LEFT)
        ttk.Combobox(mdl_frame, textvariable=self.model, values=MODELS,
                     state="readonly", width=10).pack(side=tk.LEFT)

        tk.Label(right, text=t("Regions:"), font=("Arial", 10)).pack(anchor="w", pady=(10, 2))
        self.coord_text = tk.Text(right, height=4, width=44)
        self.coord_text.pack(anchor="w", padx=5)

        btn_frame = tk.Frame(right)
        btn_frame.pack(anchor="w", padx=5, pady=8, fill=tk.X)

        tk.Button(btn_frame, text=t("Undo Last Region"),
                  command=self.undo_last_region, bg="lightblue").pack(fill=tk.X, pady=2)
        tk.Button(btn_frame, text=t("Clear All Regions (Esc)"),
                  command=self.clear_selection, bg="lightblue").pack(fill=tk.X, pady=2)
        self.process_btn = tk.Button(btn_frame, text=t("Remove Watermark"),
                                     command=self.remove_watermark, bg="lightgreen",
                                     font=("Arial", 11, "bold"))
        self.process_btn.pack(fill=tk.X, pady=4)
        tk.Button(btn_frame, text=t("Deep Clean"),
                  command=self.deep_clean, bg="#aed581").pack(fill=tk.X, pady=2)
        self.batch_btn = tk.Button(btn_frame, text=t("Batch: Apply to ALL Word Images"),
                                   command=self.batch_remove_doc, bg="#c5e1a5",
                                   font=("Arial", 10, "bold"))
        self.batch_btn.pack(fill=tk.X, pady=2)
        tk.Button(btn_frame, text=t("Reset to Original (Ctrl+Z)"),
                  command=self.reset_image, bg="lightgray").pack(fill=tk.X, pady=2)

        zoom_row = tk.Frame(btn_frame)
        zoom_row.pack(fill=tk.X, pady=2)
        tk.Button(zoom_row, text=t("Zoom -"),
                  command=lambda: self.zoom_step(1 / 1.25)).pack(
            side=tk.LEFT, expand=True, fill=tk.X)
        tk.Button(zoom_row, text=t("Fit (0)"),
                  command=self.reset_view).pack(side=tk.LEFT, expand=True, fill=tk.X)
        tk.Button(zoom_row, text=t("Zoom +"),
                  command=lambda: self.zoom_step(1.25)).pack(
            side=tk.LEFT, expand=True, fill=tk.X)

        tk.Button(btn_frame, text=t("Save to Word"), command=self.save_to_source,
                  bg="lightyellow").pack(fill=tk.X, pady=2)
        tk.Button(btn_frame, text=t("Export PNG"), command=self.save_image,
                  bg="lightcoral").pack(fill=tk.X, pady=2)

        # Progress bar
        self.progress = ttk.Progressbar(right, mode="determinate")
        self.progress.pack(fill=tk.X, padx=5, pady=(8, 2))

        tk.Label(right, text=t("Status:"), font=("Arial", 10)).pack(
            anchor="w", pady=(8, 2))
        self.status_text = tk.Text(right, height=8, width=44)
        self.status_text.pack(anchor="w", padx=5, fill=tk.BOTH, expand=True)

    def _bind_events(self):
        self.canvas.bind("<Configure>", lambda e: self.display_image())
        self.canvas.bind("<Button-1>", self.on_left_down)
        self.canvas.bind("<B1-Motion>", self.on_left_drag)
        self.canvas.bind("<ButtonRelease-1>", self.on_left_up)
        self.canvas.bind("<Double-Button-1>", self.on_left_double)
        self.canvas.bind("<Button-2>", self.on_pan_start)
        self.canvas.bind("<B2-Motion>", self.on_pan_move)
        self.canvas.bind("<ButtonRelease-2>", self.on_pan_end)
        self.canvas.bind("<MouseWheel>", self.on_wheel)
        self.canvas.bind("<Button-4>", lambda e: self._zoom_around(e.x, e.y, 1.1))
        self.canvas.bind("<Button-5>", lambda e: self._zoom_around(e.x, e.y, 1 / 1.1))
        self.canvas.bind("<Motion>", self._on_cursor_motion)
        self.canvas.bind("<Leave>", lambda e: self.canvas.delete("cursor"))

        self.root.bind("<Escape>", lambda e: self.clear_selection())
        self.root.bind("<Control-z>", lambda e: self.undo())
        self.root.bind("<Control-y>", lambda e: self.redo())
        self.root.bind("<Control-Z>", lambda e: self.undo())
        self.root.bind("<Control-Y>", lambda e: self.redo())
        self.root.bind("<plus>", lambda e: self.zoom_step(1.25))
        self.root.bind("<equal>", lambda e: self.zoom_step(1.25))
        self.root.bind("<minus>", lambda e: self.zoom_step(1 / 1.25))
        self.root.bind("<Key-0>", lambda e: self.reset_view())
        self.root.bind("<Return>", lambda e: self.finalize_polygon())
        # Tool mode hotkeys
        self.root.bind("<Key-r>", lambda e: self._set_tool("rect"))
        self.root.bind("<Key-R>", lambda e: self._set_tool("rect"))
        self.root.bind("<Key-b>", lambda e: self._set_tool("brush"))
        self.root.bind("<Key-B>", lambda e: self._set_tool("brush"))
        self.root.bind("<Key-p>", lambda e: self._set_tool("polygon"))
        self.root.bind("<Key-P>", lambda e: self._set_tool("polygon"))
        self.root.bind("<Key-c>", lambda e: self._set_tool("color"))
        self.root.bind("<Key-C>", lambda e: self._set_tool("color"))
        self.root.bind("<Key-l>", lambda e: self.toggle_loupe())
        self.root.bind("<Key-L>", lambda e: self.toggle_loupe())
        self.root.bind("<Delete>", lambda e: self._delete_selected_rect())
        self.root.bind("<BackSpace>", lambda e: self._delete_selected_rect())

    def _setup_dnd(self):
        if not HAVE_DND:
            return
        try:
            self.root.drop_target_register(DND_FILES)
            self.root.dnd_bind("<<Drop>>", self._on_drop)
        except Exception:
            pass

    def _on_drop(self, event):
        # event.data is space-separated {path1} {path2} on Windows.
        data = event.data or ""
        paths = []
        buf = ""
        in_brace = False
        for ch in data:
            if ch == "{":
                in_brace = True
                buf = ""
            elif ch == "}":
                in_brace = False
                paths.append(buf)
                buf = ""
            elif ch == " " and not in_brace:
                if buf:
                    paths.append(buf)
                buf = ""
            else:
                buf += ch
        if buf:
            paths.append(buf)
        if not paths:
            return
        p = paths[0]
        self._open_by_extension(p)

    def _open_by_extension(self, p):
        ext = Path(p).suffix.lower()
        if ext in IMAGE_EXTS:
            self._load_image_file(p)
        elif ext == ".docx":
            self._load_ooxml(p, "word")
        elif ext == ".pptx":
            self._load_ooxml(p, "pptx")
        elif ext == ".xlsx":
            self._load_ooxml(p, "xlsx")
        elif ext == ".pdf":
            self._load_pdf(p)
        elif ext == TEMPLATE_EXT:
            self._apply_template_file(p)
        else:
            self.log("Unsupported file: " + p)

    def log(self, message):
        self.status_text.insert(tk.END, message + "\n")
        self.status_text.see(tk.END)
        self.root.update()

    def _refresh_file_text(self):
        self.file_text.delete(1.0, tk.END)
        if self._file_desc is None:
            return
        kind = self._file_desc[0]
        if kind == "image":
            self.file_text.insert(tk.END, t("Type: Image file\n") + self._file_desc[1])
        else:
            _, doc_kind, name, extra = self._file_desc
            self.file_text.insert(
                tk.END,
                t("Type: {} document\n{}\nrId: {}").format(doc_kind, name, extra),
            )

    def set_language(self, lang):
        global LANG
        if lang == LANG:
            return
        LANG = lang
        CONFIG["language"] = lang
        save_config(CONFIG)
        for child in self.root.winfo_children():
            child.destroy()
        self.root.config(menu=tk.Menu(self.root))
        self._build_menu()
        self._build_layout()
        self._bind_events()
        self.root.title(t("Watermark Remover v5"))
        self._refresh_file_text()
        if self.cv_image is not None:
            self.reset_view()
        self._update_coord_text()
        self.log(t("Language switched to: ") + ("中文" if lang == "zh" else "English"))

    # ------------------------- recent
    def _remember(self, path):
        recents = CONFIG.get("recent", [])
        path = str(path)
        if path in recents:
            recents.remove(path)
        recents.insert(0, path)
        CONFIG["recent"] = recents[:10]
        CONFIG["last_dir"] = str(Path(path).parent)
        save_config(CONFIG)
        self._rebuild_recent_menu()

    def clear_recent(self):
        CONFIG["recent"] = []
        save_config(CONFIG)
        self._rebuild_recent_menu()

    def open_recent(self, path):
        self._open_by_extension(path)

    # ------------------------- image load
    def _load_image_file(self, filepath):
        try:
            img = cv2.imdecode(np.fromfile(filepath, dtype=np.uint8), cv2.IMREAD_COLOR)
        except Exception as e:
            self.log(t("Read error: ") + str(e))
            img = None
        if img is None:
            messagebox.showerror(t("Error"), t("Cannot open image:\n") + filepath)
            return
        self.cv_image = img
        self.original_image = img.copy()
        self.source_kind = "image"
        self.source_path = filepath
        self.source_extra = None
        self.word_doc = None
        self.pdf_doc = None
        self._file_desc = ("image", Path(filepath).name)
        self._reset_selection_state()
        self._reset_undo()
        self._refresh_file_text()
        self.log(t("Opened: ") + Path(filepath).name)
        self.log(t("Size: {}x{}").format(img.shape[1], img.shape[0]))
        self.reset_view()
        self._remember(filepath)

    def open_image(self):
        initial = CONFIG.get("last_dir") or None
        filepath = filedialog.askopenfilename(
            initialdir=initial,
            filetypes=[("Image files", "*.png *.jpg *.jpeg *.bmp *.webp *.tif *.tiff"),
                       ("All files", "*.*")],
        )
        if filepath:
            self._load_image_file(filepath)

    # ------------------------- OOXML (docx/pptx/xlsx)
    def open_from_ooxml(self, kind):
        ext = {"word": "*.docx", "pptx": "*.pptx", "xlsx": "*.xlsx"}[kind]
        initial = CONFIG.get("last_dir") or None
        filepath = filedialog.askopenfilename(
            initialdir=initial,
            filetypes=[(kind + " files", ext), ("All files", "*.*")],
        )
        if filepath:
            self._load_ooxml(filepath, kind)

    def _load_ooxml(self, filepath, kind):
        try:
            images_info = []
            for partname, data in list_ooxml_images(filepath):
                images_info.append({"partname": partname,
                                    "name": Path(partname).name,
                                    "data": data,
                                    "size": len(data)})
        except Exception as e:
            messagebox.showerror(t("Error"), t("Cannot open archive: ") + str(e))
            return
        if not images_info:
            messagebox.showerror(t("Error"), t("No images found"))
            return
        self.source_kind = kind
        self.source_path = filepath
        self.word_doc = None
        self.pdf_doc = None
        if kind == "word":
            try:
                self.word_doc = Document(filepath)
            except Exception:
                pass
        self.log(t("Opened: ") + Path(filepath).name)
        self.log(t("Found {} images").format(len(images_info)))
        self._remember(filepath)
        self._show_image_selector(images_info, source="ooxml")

    # ------------------------- PDF
    def open_from_pdf(self):
        if not HAVE_PDF:
            messagebox.showwarning(
                t("Warning"), t("PyMuPDF not installed. Run: pip install pymupdf"))
            return
        initial = CONFIG.get("last_dir") or None
        filepath = filedialog.askopenfilename(
            initialdir=initial,
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
        )
        if filepath:
            self._load_pdf(filepath)

    def _load_pdf(self, filepath):
        if not HAVE_PDF:
            messagebox.showwarning(
                t("Warning"), t("PyMuPDF not installed. Run: pip install pymupdf"))
            return
        try:
            doc = fitz.open(filepath)
        except Exception as e:
            messagebox.showerror(t("Error"), t("Cannot open PDF: ") + str(e))
            return
        images_info = []
        seen = set()
        for page_index in range(len(doc)):
            page = doc[page_index]
            for img_info in page.get_images(full=True):
                xref = img_info[0]
                if xref in seen:
                    continue
                seen.add(xref)
                try:
                    data = doc.extract_image(xref)
                    images_info.append({
                        "partname": "pdf:xref{}".format(xref),
                        "name": "page{}_img{}.{}".format(
                            page_index + 1, xref, data.get("ext", "png")),
                        "data": data["image"],
                        "size": len(data["image"]),
                        "xref": xref,
                    })
                except Exception:
                    continue
        if not images_info:
            doc.close()
            messagebox.showerror(t("Error"), t("No images found"))
            return
        self.source_kind = "pdf"
        self.source_path = filepath
        self.pdf_doc = doc
        self.word_doc = None
        self.log(t("Opened: ") + Path(filepath).name)
        self.log(t("Found {} images").format(len(images_info)))
        self._remember(filepath)
        self._show_image_selector(images_info, source="pdf")

    # ------------------------- image-selector dialog
    def _show_image_selector(self, images_info, source):
        selector = tk.Toplevel(self.root)
        selector.title(t("Select Image"))
        selector.geometry("720x520")
        tk.Label(selector, text=t("Select image to edit:"),
                 font=("Arial", 11, "bold")).pack(pady=10)

        main_frame = tk.Frame(selector)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        list_frame = tk.Frame(main_frame)
        list_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tk.Label(list_frame, text=t("Images:"), font=("Arial", 10, "bold")).pack(anchor="w")
        frame = tk.Frame(list_frame)
        frame.pack(fill=tk.BOTH, expand=True)
        scrollbar = tk.Scrollbar(frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        listbox = tk.Listbox(frame, yscrollcommand=scrollbar.set, height=12)
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=listbox.yview)
        for info in images_info:
            listbox.insert(tk.END, "{} ({:.2f} MB)".format(
                info["name"], info["size"] / 1024 / 1024))

        preview_frame = tk.Frame(main_frame, width=280, bg="gray")
        preview_frame.pack(side=tk.RIGHT, fill=tk.BOTH, padx=(10, 0))
        tk.Label(preview_frame, text=t("Preview:"),
                 font=("Arial", 10, "bold"), bg="gray", fg="white").pack()
        preview_canvas = tk.Canvas(preview_frame, bg="gray", width=260, height=360)
        preview_canvas.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        def _preview(_=None):
            sel = listbox.curselection()
            if not sel:
                return
            info = images_info[sel[0]]
            nparr = np.frombuffer(info["data"], np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            if img is None:
                return
            h, w = img.shape[:2]
            s = min(260 / w, 360 / h, 1.0)
            nw, nh = max(1, int(w * s)), max(1, int(h * s))
            rgb = cv2.cvtColor(cv2.resize(img, (nw, nh)), cv2.COLOR_BGR2RGB)
            photo = ImageTk.PhotoImage(Image.fromarray(rgb))
            preview_canvas.delete("all")
            preview_canvas.create_image(130, 180, image=photo)
            preview_canvas.image = photo

        listbox.bind("<<ListboxSelect>>", _preview)
        if images_info:
            listbox.selection_set(0)
            _preview()

        def _confirm():
            sel = listbox.curselection()
            if not sel:
                messagebox.showwarning(t("Warning"), t("Please select an image"))
                return
            info = images_info[sel[0]]
            nparr = np.frombuffer(info["data"], np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            if img is None:
                messagebox.showerror(t("Error"), t("Cannot decode image"))
                return
            self.cv_image = img
            self.original_image = img.copy()
            if source == "pdf":
                self.source_extra = info["xref"]
                kind_label = "PDF"
            else:
                self.source_extra = info["partname"]
                kind_label = self.source_kind.upper()
            self._file_desc = ("doc", kind_label, info["name"], str(self.source_extra))
            self._reset_selection_state()
            self._reset_undo()
            self._refresh_file_text()
            self.log(t("Loaded: ") + info["name"])
            self.log(t("Size: {}x{}").format(img.shape[1], img.shape[0]))
            self.reset_view()
            selector.destroy()

        tk.Button(selector, text=t("Confirm"), command=_confirm,
                  bg="lightgreen", font=("Arial", 11)).pack(pady=10)

    # -------------------- view
    def reset_view(self):
        if self.cv_image is None:
            self.zoom = 1.0
            self.pan_x = 0.0
            self.pan_y = 0.0
            self.canvas.delete("all")
            return
        canvas_w = self.canvas.winfo_width()
        canvas_h = self.canvas.winfo_height()
        if canvas_w < 50 or canvas_h < 50:
            canvas_w, canvas_h = 800, 600
        h, w = self.cv_image.shape[:2]
        self.base_scale = min(canvas_w / w, canvas_h / h, 1.0)
        self.zoom = 1.0
        eff = self.base_scale * self.zoom
        self.pan_x = (canvas_w - w * eff) / 2
        self.pan_y = (canvas_h - h * eff) / 2
        self.display_image()

    def zoom_step(self, factor):
        if self.cv_image is None:
            return
        cx = self.canvas.winfo_width() / 2
        cy = self.canvas.winfo_height() / 2
        self._zoom_around(cx, cy, factor)

    def _zoom_around(self, cx, cy, factor):
        if self.cv_image is None:
            return
        new_zoom = max(0.1, min(20.0, self.zoom * factor))
        if abs(new_zoom - self.zoom) < 1e-6:
            return
        ix, iy = self._canvas_to_image(cx, cy)
        self.zoom = new_zoom
        eff = self.base_scale * self.zoom
        self.pan_x = cx - ix * eff
        self.pan_y = cy - iy * eff
        self.display_image()

    def on_wheel(self, event):
        if self.cv_image is None:
            return
        factor = 1.1 if event.delta > 0 else (1 / 1.1)
        self._zoom_around(event.x, event.y, factor)

    def display_image(self):
        if self.cv_image is None:
            return
        canvas_w = self.canvas.winfo_width()
        canvas_h = self.canvas.winfo_height()
        if canvas_w < 50 or canvas_h < 50:
            canvas_w, canvas_h = 800, 600
        h, w = self.cv_image.shape[:2]
        self.base_scale = min(canvas_w / w, canvas_h / h, 1.0)
        eff = self.base_scale * self.zoom
        nw = max(1, int(w * eff))
        nh = max(1, int(h * eff))
        if nw * nh > 40_000_000:
            cap = (40_000_000 / (nw * nh)) ** 0.5
            nw = max(1, int(nw * cap))
            nh = max(1, int(nh * cap))
        rgb = cv2.cvtColor(self.cv_image, cv2.COLOR_BGR2RGB)
        # Overlay extra_mask in translucent red before scaling.
        if self.extra_mask is not None and self.extra_mask.max() > 0:
            overlay = rgb.copy()
            red = np.zeros_like(rgb)
            red[..., 0] = 255
            alpha = (self.extra_mask.astype(np.float32) / 255.0 * 0.4)[..., None]
            rgb = (overlay * (1 - alpha) + red * alpha).astype(np.uint8)
        resized = cv2.resize(rgb, (nw, nh))
        self.photo = ImageTk.PhotoImage(Image.fromarray(resized))
        self.canvas.delete("all")
        self.canvas.create_image(self.pan_x, self.pan_y, anchor=tk.NW,
                                 image=self.photo, tags="img")
        self._redraw_overlay()

    def _image_to_canvas(self, ix, iy):
        eff = self.base_scale * self.zoom
        return self.pan_x + ix * eff, self.pan_y + iy * eff

    def _canvas_to_image(self, cx, cy):
        eff = self.base_scale * self.zoom
        if eff <= 0 or self.cv_image is None:
            return 0, 0
        x = (cx - self.pan_x) / eff
        y = (cy - self.pan_y) / eff
        h, w = self.cv_image.shape[:2]
        return max(0, min(w, x)), max(0, min(h, y))

    def _redraw_overlay(self):
        self.canvas.delete("rect")
        for idx, (x1, y1, x2, y2) in enumerate(self.rects):
            cx1, cy1 = self._image_to_canvas(x1, y1)
            cx2, cy2 = self._image_to_canvas(x2, y2)
            is_sel = idx == self.selected_rect_idx
            color = "#00ffff" if is_sel else "red"
            width = 3 if is_sel else 2
            self.canvas.create_rectangle(cx1, cy1, cx2, cy2,
                                         outline=color, width=width, tags="rect")
            if is_sel:
                for hx, hy in ((cx1, cy1), (cx2, cy1), (cx1, cy2), (cx2, cy2)):
                    self.canvas.create_rectangle(hx - 4, hy - 4, hx + 4, hy + 4,
                                                 outline=color, fill=color,
                                                 tags="rect")
        if self.in_progress:
            x1, y1, x2, y2 = self.in_progress
            cx1, cy1 = self._image_to_canvas(x1, y1)
            cx2, cy2 = self._image_to_canvas(x2, y2)
            self.canvas.create_rectangle(cx1, cy1, cx2, cy2,
                                         outline="orange", width=2,
                                         dash=(4, 2), tags="rect")
        if self.polygon_points:
            pts_canvas = []
            for ix, iy in self.polygon_points:
                cx, cy = self._image_to_canvas(ix, iy)
                pts_canvas.extend([cx, cy])
            if len(self.polygon_points) >= 2:
                self.canvas.create_line(*pts_canvas, fill="orange", width=2,
                                        dash=(4, 2), tags="rect")
            for ix, iy in self.polygon_points:
                cx, cy = self._image_to_canvas(ix, iy)
                self.canvas.create_oval(cx - 3, cy - 3, cx + 3, cy + 3,
                                        fill="orange", outline="red", tags="rect")

    # -------------------- mouse
    def _set_tool(self, name):
        if name in ("rect", "brush", "polygon", "color"):
            self.tool_mode.set(name)

    def _rect_hit_test(self, ix, iy):
        """Return index of rect under (ix, iy), or None."""
        for idx in range(len(self.rects) - 1, -1, -1):  # top-most first
            x1, y1, x2, y2 = self.rects[idx]
            if min(x1, x2) <= ix <= max(x1, x2) and min(y1, y2) <= iy <= max(y1, y2):
                return idx
        return None

    def _delete_selected_rect(self):
        if self.selected_rect_idx is None:
            return
        if 0 <= self.selected_rect_idx < len(self.rects):
            self.rects.pop(self.selected_rect_idx)
            self.selected_rect_idx = None
            self._redraw_overlay()
            self._update_coord_text()
            self.log(t("Deleted region"))

    def on_left_down(self, event):
        if self.cv_image is None:
            return
        if event.state & self.SHIFT_MASK:
            self.on_pan_start(event)
            return
        if self._color_pick_pending:
            self._apply_color_pick(event.x, event.y)
            return

        mode = self.tool_mode.get()
        ix, iy = self._canvas_to_image(event.x, event.y)

        if mode == "rect":
            # If we click inside an existing rect, select it and start a move.
            hit = self._rect_hit_test(ix, iy)
            if hit is not None:
                self.selected_rect_idx = hit
                self.rect_drag = (tuple(self.rects[hit]), ix, iy)
                self.drawing = True
                self._redraw_overlay()
                self.log(t("Selected region #{}").format(hit + 1))
                return
            # Clicking empty space deselects and starts a new rect.
            self.selected_rect_idx = None
            self.rect_drag = None
            self.in_progress = (ix, iy, ix, iy)
            self.drawing = True
            self._redraw_overlay()
        elif mode == "brush":
            self._ensure_extra_mask()
            self.drawing = True
            self._paint_brush(ix, iy)
            self.display_image()
        elif mode == "polygon":
            self.polygon_points.append((ix, iy))
            self._redraw_overlay()
        elif mode == "color":
            self._apply_color_pick_at_img(ix, iy)

    def on_left_drag(self, event):
        if self.panning:
            self.on_pan_move(event)
            return
        if not self.drawing or self.cv_image is None:
            return
        ix, iy = self._canvas_to_image(event.x, event.y)
        mode = self.tool_mode.get()
        if mode == "rect":
            if self.rect_drag is not None:
                # Moving an existing rect.
                (ox1, oy1, ox2, oy2), sx, sy = self.rect_drag
                dx = ix - sx
                dy = iy - sy
                self.rects[self.selected_rect_idx] = (ox1 + dx, oy1 + dy,
                                                      ox2 + dx, oy2 + dy)
                self._redraw_overlay()
                self._update_coord_text()
            elif self.in_progress is not None:
                x1, y1, _, _ = self.in_progress
                self.in_progress = (x1, y1, ix, iy)
                self._redraw_overlay()
                self._update_coord_text()
        elif mode == "brush":
            self._paint_brush(ix, iy)
            self.display_image()

    def on_left_up(self, event):
        if self.panning:
            self.on_pan_end(event)
            return
        if not self.drawing:
            return
        self.drawing = False
        mode = self.tool_mode.get()
        if mode == "rect":
            if self.rect_drag is not None:
                self.rect_drag = None
                self._redraw_overlay()
                self._update_coord_text()
            elif self.in_progress:
                x1, y1, x2, y2 = self.in_progress
                x1, x2 = min(x1, x2), max(x1, x2)
                y1, y2 = min(y1, y2), max(y1, y2)
                if (x2 - x1) >= 2 and (y2 - y1) >= 2:
                    self.rects.append((x1, y1, x2, y2))
                self.in_progress = None
                self._redraw_overlay()
                self._update_coord_text()

    def on_left_double(self, event):
        if self.tool_mode.get() == "polygon":
            self.finalize_polygon()

    def finalize_polygon(self):
        if self.tool_mode.get() != "polygon" or len(self.polygon_points) < 3:
            return
        self._ensure_extra_mask()
        pts = np.array([[int(x), int(y)] for x, y in self.polygon_points], dtype=np.int32)
        cv2.fillPoly(self.extra_mask, [pts], 255)
        self.polygon_points = []
        self.display_image()
        self._update_coord_text()

    def _on_cursor_motion(self, event):
        self.canvas.delete("cursor")
        if self.cv_image is None or self.panning:
            return
        cw = self.canvas.winfo_width()
        ch = self.canvas.winfo_height()
        # High-contrast crosshair so users always see the exact pointer
        # position even over a busy image.
        self.canvas.create_line(0, event.y, cw, event.y,
                                fill="yellow", dash=(3, 3), tags="cursor")
        self.canvas.create_line(event.x, 0, event.x, ch,
                                fill="yellow", dash=(3, 3), tags="cursor")
        if self.tool_mode.get() == "brush":
            r = max(1, int(self.brush_size.get())) * self.base_scale * self.zoom
            self.canvas.create_oval(
                event.x - r, event.y - r, event.x + r, event.y + r,
                outline="yellow", width=1, tags="cursor")
        if self.loupe_win is not None:
            ix, iy = self._canvas_to_image(event.x, event.y)
            self._update_loupe(ix, iy)

    def on_pan_start(self, event):
        if self.cv_image is None:
            return
        self.panning = True
        self.pan_anchor = (event.x, event.y, self.pan_x, self.pan_y)
        self.canvas.config(cursor="fleur")
        self.canvas.delete("cursor")

    def on_pan_move(self, event):
        if not self.panning or not self.pan_anchor:
            return
        ex0, ey0, px0, py0 = self.pan_anchor
        self.pan_x = px0 + (event.x - ex0)
        self.pan_y = py0 + (event.y - ey0)
        self.display_image()

    def on_pan_end(self, event):
        self.panning = False
        self.pan_anchor = None
        self.canvas.config(cursor="tcross")

    # -------------------- brush / color / poly helpers
    def _ensure_extra_mask(self):
        if self.extra_mask is None and self.cv_image is not None:
            h, w = self.cv_image.shape[:2]
            self.extra_mask = np.zeros((h, w), dtype=np.uint8)

    def _paint_brush(self, ix, iy):
        if self.extra_mask is None:
            return
        r = max(1, int(self.brush_size.get()))
        cv2.circle(self.extra_mask, (int(ix), int(iy)), r, 255, -1)

    def begin_color_pick(self):
        self._color_pick_pending = True
        self.log(t("Click a pixel to sample color"))

    def _apply_color_pick(self, cx, cy):
        ix, iy = self._canvas_to_image(cx, cy)
        self._apply_color_pick_at_img(ix, iy)

    def _apply_color_pick_at_img(self, ix, iy):
        if self.cv_image is None:
            return
        h, w = self.cv_image.shape[:2]
        x = max(0, min(w - 1, int(ix)))
        y = max(0, min(h - 1, int(iy)))
        target = self.cv_image[y, x].astype(np.int16)
        tol = int(self.color_tol.get())
        diff = np.abs(self.cv_image.astype(np.int16) - target).max(axis=2)
        sel = (diff <= tol).astype(np.uint8) * 255
        self._ensure_extra_mask()
        self.extra_mask = cv2.bitwise_or(self.extra_mask, sel)
        self._color_pick_pending = False
        self.display_image()
        self._update_coord_text()

    # -------------------- selection utilities
    def undo_last_region(self):
        if self.rects:
            self.rects.pop()
            self._redraw_overlay()
            self._update_coord_text()
            self.log(t("Removed last region"))

    def clear_selection(self):
        self.rects = []
        self.in_progress = None
        self.drawing = False
        self.extra_mask = None
        self.polygon_points = []
        self.display_image()
        self._update_coord_text()
        self.log(t("Cleared all regions"))

    def _reset_selection_state(self):
        self.rects = []
        self.in_progress = None
        self.drawing = False
        self.extra_mask = None
        self.polygon_points = []
        self.coord_text.delete(1.0, tk.END)

    def _update_coord_text(self):
        self.coord_text.delete(1.0, tk.END)
        text = t("Total regions: {}\n").format(len(self.rects))
        start = max(0, len(self.rects) - 3)
        for i in range(start, len(self.rects)):
            x1, y1, x2, y2 = self.rects[i]
            text += "  #{}: x=[{}:{}] y=[{}:{}] {}x{}\n".format(
                i + 1, int(x1), int(x2), int(y1), int(y2),
                int(x2 - x1), int(y2 - y1))
        if self.extra_mask is not None and self.extra_mask.max() > 0:
            nonzero = int((self.extra_mask > 0).sum())
            text += "mask px: {}\n".format(nonzero)
        if self.in_progress:
            x1, y1, x2, y2 = self.in_progress
            x1, x2 = min(x1, x2), max(x1, x2)
            y1, y2 = min(y1, y2), max(y1, y2)
            text += t("drawing: x=[{}:{}] y=[{}:{}]\n").format(
                int(x1), int(x2), int(y1), int(y2))
        self.coord_text.insert(tk.END, text)

    def reset_image(self):
        if self.original_image is None:
            return
        self._push_undo()
        self.cv_image = self.original_image.copy()
        self._reset_selection_state()
        self.reset_view()
        self.log(t("Reset to original"))

    # -------------------- undo / redo
    def _reset_undo(self):
        self.undo_stack = []
        self.redo_stack = []

    def _snapshot(self):
        return (self.cv_image.copy() if self.cv_image is not None else None,
                list(self.rects),
                self.extra_mask.copy() if self.extra_mask is not None else None)

    def _push_undo(self):
        self.undo_stack.append(self._snapshot())
        if len(self.undo_stack) > 20:
            self.undo_stack.pop(0)
        self.redo_stack.clear()

    def undo(self):
        if not self.undo_stack:
            self.log(t("Nothing to undo"))
            return
        self.redo_stack.append(self._snapshot())
        img, rects, mask = self.undo_stack.pop()
        self.cv_image = img
        self.rects = list(rects)
        self.extra_mask = mask
        self._update_coord_text()
        self.display_image()
        self.log(t("Undone"))

    def redo(self):
        if not self.redo_stack:
            self.log(t("Nothing to redo"))
            return
        self.undo_stack.append(self._snapshot())
        img, rects, mask = self.redo_stack.pop()
        self.cv_image = img
        self.rects = list(rects)
        self.extra_mask = mask
        self._update_coord_text()
        self.display_image()
        self.log(t("Redone"))

    # -------------------- mask composition + inpaint
    def _build_combined_mask(self, h, w, rects=None, extra_mask=None, feather=None):
        mask = np.zeros((h, w), dtype=np.uint8)
        if rects is None:
            rects = self.rects
        for (x1, y1, x2, y2) in rects:
            ix1 = max(0, int(min(x1, x2)))
            iy1 = max(0, int(min(y1, y2)))
            ix2 = min(w, int(max(x1, x2)))
            iy2 = min(h, int(max(y1, y2)))
            if ix2 > ix1 and iy2 > iy1:
                mask[iy1:iy2, ix1:ix2] = 255
        if extra_mask is None:
            extra_mask = self.extra_mask
        if extra_mask is not None and extra_mask.shape == mask.shape:
            mask = cv2.bitwise_or(mask, extra_mask)
        f = self.feather.get() if feather is None else feather
        if f > 0:
            k = int(f) * 2 + 1
            mask = cv2.GaussianBlur(mask, (k, k), 0)
            _, mask = cv2.threshold(mask, 20, 255, cv2.THRESH_BINARY)
        return mask

    def _run_iopaint(self, img_bgr, mask, temp_dir):
        temp_dir.mkdir(exist_ok=True, parents=True)
        img_path = temp_dir / "input.png"
        mask_path = temp_dir / "mask.png"
        out_dir = temp_dir / "output"
        if out_dir.exists():
            shutil.rmtree(out_dir)
        out_dir.mkdir(parents=True)
        cv2.imwrite(str(img_path), img_bgr)
        cv2.imwrite(str(mask_path), mask)
        model = self.model.get() or "lama"
        if getattr(sys, "frozen", False):
            cmd = [sys.executable, "--iopaint-worker", "run",
                   "--model", model,
                   "--image", str(img_path),
                   "--mask", str(mask_path),
                   "--output", str(out_dir)]
        else:
            cmd = [sys.executable, "-m", "iopaint", "run",
                   "--model", model,
                   "--image", str(img_path),
                   "--mask", str(mask_path),
                   "--output", str(out_dir)]
        subprocess.run(cmd, check=True, capture_output=True, timeout=600,
                       text=True, encoding="utf-8", errors="replace")
        result = cv2.imread(str(out_dir / "input.png"))
        if result is None:
            raise RuntimeError("iopaint produced no output at " + str(out_dir / "input.png"))
        return result

    def _start_indeterminate(self):
        try:
            self.progress.configure(mode="indeterminate")
            self.progress.start(80)
        except Exception:
            pass

    def _stop_progress(self):
        try:
            self.progress.stop()
            self.progress.configure(mode="determinate", value=0)
        except Exception:
            pass

    def remove_watermark(self):
        self._remove_core(show_preview=True, show_success=True)

    def _remove_core(self, show_preview=True, show_success=False):
        if self.cv_image is None:
            messagebox.showerror(t("Error"), t("Please open an image first"))
            return False
        if not self.rects and (self.extra_mask is None or self.extra_mask.max() == 0):
            messagebox.showerror(t("Error"), t("Draw at least one region first"))
            return False
        h, w = self.cv_image.shape[:2]
        mask = self._build_combined_mask(h, w)
        if mask.max() == 0:
            messagebox.showerror(t("Error"), t("Selection is empty"))
            return False
        self.log(t("Removing watermark..."))
        self.process_btn.config(state="disabled")
        self.batch_btn.config(state="disabled")
        self.root.config(cursor="watch")
        self._start_indeterminate()
        self.root.update()
        try:
            result = self._run_iopaint(self.cv_image, mask, TEMP_ROOT)
        except subprocess.CalledProcessError as e:
            err = (e.stderr or e.stdout or "")[-800:]
            self.log(t("iopaint failed:\n") + err)
            messagebox.showerror(t("Error"), t("iopaint failed:\n") + err)
            self._stop_progress()
            self.process_btn.config(state="normal")
            self.batch_btn.config(state="normal")
            self.root.config(cursor="")
            return False
        except subprocess.TimeoutExpired:
            self.log(t("iopaint timed out (>10min)"))
            messagebox.showerror(t("Error"), t("iopaint timed out (>10min)"))
            self._stop_progress()
            self.process_btn.config(state="normal")
            self.batch_btn.config(state="normal")
            self.root.config(cursor="")
            return False
        except Exception as e:
            self.log(t("Failed: ") + str(e))
            messagebox.showerror(t("Error"), t("Failed: ") + str(e))
            self._stop_progress()
            self.process_btn.config(state="normal")
            self.batch_btn.config(state="normal")
            self.root.config(cursor="")
            return False

        self._stop_progress()
        self.process_btn.config(state="normal")
        self.batch_btn.config(state="normal")
        self.root.config(cursor="")

        if show_preview:
            if not self._preview_remove_result(result):
                self.log(t("Cancel"))
                return False

        self._push_undo()
        self.cv_image = result
        self.rects = []
        self.in_progress = None
        self.extra_mask = None
        self.polygon_points = []
        self.selected_rect_idx = None
        self._update_coord_text()
        self.display_image()
        self.log(t("Done."))
        if show_success:
            messagebox.showinfo(t("Success"), t("Watermark removed"))
        return True

    # -------------------- batch modes
    def batch_remove_doc(self):
        if self.source_kind not in ("word", "pptx", "xlsx", "pdf"):
            messagebox.showerror(
                t("Error"), t("Open a Word document first (File -> Open from Word)"))
            return
        if self.original_image is None or (not self.rects and (self.extra_mask is None or self.extra_mask.max() == 0)):
            messagebox.showerror(
                t("Error"), t("Draw at least one region on the current image first"))
            return
        h0, w0 = self.original_image.shape[:2]
        rel_rects = [(x1 / w0, y1 / h0, x2 / w0, y2 / h0)
                     for (x1, y1, x2, y2) in self.rects]
        extra_ref = self.extra_mask.copy() if self.extra_mask is not None else None

        if self.source_kind == "pdf":
            self._batch_pdf(rel_rects, extra_ref, h0, w0)
        else:
            self._batch_ooxml(rel_rects, extra_ref, h0, w0)

    def _scale_mask_to(self, extra_ref, h0, w0, h, w):
        if extra_ref is None:
            return None
        if extra_ref.shape == (h, w):
            return extra_ref
        return cv2.resize(extra_ref, (w, h), interpolation=cv2.INTER_NEAREST)

    def _batch_ooxml(self, rel_rects, extra_ref, h0, w0):
        try:
            images = list_ooxml_images(self.source_path)
        except Exception as e:
            messagebox.showerror(t("Error"), t("Cannot open archive: ") + str(e))
            return
        if not messagebox.askyesno(
            t("Batch Remove"),
            t("Apply the selection you drew to ALL {} images in:\n"
              "{}\n\n"
              "Mask scales proportionally per image.\n"
              "A .backup copy of the file will be created.\n\nContinue?").format(
                  len(images), self.source_path)):
            return

        backup_path = str(self.source_path) + ".backup"
        shutil.copy(self.source_path, backup_path)
        self.log(t("Backup: ") + backup_path)

        self.process_btn.config(state="disabled")
        self.batch_btn.config(state="disabled")
        self.root.config(cursor="watch")
        self.progress.configure(mode="determinate", maximum=len(images), value=0)
        self.root.update()
        replacements = {}
        failed = []
        try:
            for idx, (partname, data) in enumerate(images, 1):
                name = Path(partname).name
                self.log(t("[{}/{}] {}").format(idx, len(images), name))
                self.progress["value"] = idx - 1
                self.root.update()
                try:
                    img = cv2.imdecode(np.frombuffer(data, np.uint8), cv2.IMREAD_COLOR)
                    if img is None:
                        self.log(t("  skip (cannot decode)"))
                        failed.append(name)
                        continue
                    h, w = img.shape[:2]
                    abs_rects = [(rx1 * w, ry1 * h, rx2 * w, ry2 * h)
                                 for (rx1, ry1, rx2, ry2) in rel_rects]
                    extra = self._scale_mask_to(extra_ref, h0, w0, h, w)
                    mask = self._build_combined_mask(h, w, rects=abs_rects,
                                                     extra_mask=extra)
                    if mask.max() == 0:
                        self.log(t("  skip (empty mask)"))
                        failed.append(name)
                        continue
                    result = self._run_iopaint(img, mask, BATCH_TEMP_ROOT)
                    ext = Path(partname).suffix.lower() or ".png"
                    ok, buf = cv2.imencode(ext, result)
                    if not ok:
                        self.log(t("  encode failed"))
                        failed.append(name)
                        continue
                    replacements[partname] = buf.tobytes()
                    self.log(t("  ok"))
                except subprocess.CalledProcessError as e:
                    err = (e.stderr or e.stdout or "")[-300:]
                    self.log(t("  FAILED: iopaint: ") + err)
                    failed.append(name)
                except Exception as e:
                    self.log(t("  FAILED: ") + str(e))
                    failed.append(name)
            self.progress["value"] = len(images)
            if not replacements:
                messagebox.showerror(t("Batch Remove"), t("No images were processed."))
                return
            replace_ooxml_images(self.source_path, replacements, BATCH_EXTRACT_ROOT)
            if self.source_kind == "word":
                try:
                    self.word_doc = Document(self.source_path)
                except Exception:
                    pass
            msg = t("Batch complete.\nProcessed: {}\nFailed: {}\nBackup: {}").format(
                len(replacements), len(failed), backup_path)
            if failed:
                msg += t("\n\nFailed images:\n") + "\n".join("  " + f for f in failed[:10])
                if len(failed) > 10:
                    msg += t("\n  ... and {} more").format(len(failed) - 10)
            self.log(msg)
            messagebox.showinfo(t("Batch Remove"), msg)
        finally:
            self._stop_progress()
            self.process_btn.config(state="normal")
            self.batch_btn.config(state="normal")
            self.root.config(cursor="")

    def _batch_pdf(self, rel_rects, extra_ref, h0, w0):
        if not HAVE_PDF or self.pdf_doc is None:
            messagebox.showwarning(
                t("Warning"), t("PyMuPDF not installed. Run: pip install pymupdf"))
            return
        doc = self.pdf_doc
        all_xrefs = []
        seen = set()
        for page in doc:
            for info in page.get_images(full=True):
                xref = info[0]
                if xref not in seen:
                    seen.add(xref)
                    all_xrefs.append(xref)
        if not all_xrefs:
            messagebox.showinfo(t("Info"), t("No images found"))
            return
        if not messagebox.askyesno(
            t("Batch Remove"),
            t("Apply the selection you drew to ALL {} images in:\n"
              "{}\n\n"
              "Mask scales proportionally per image.\n"
              "A .backup copy of the file will be created.\n\nContinue?").format(
                  len(all_xrefs), self.source_path)):
            return

        backup_path = str(self.source_path) + ".backup"
        doc.save(backup_path)
        self.log(t("Backup: ") + backup_path)

        self.process_btn.config(state="disabled")
        self.batch_btn.config(state="disabled")
        self.root.config(cursor="watch")
        self.progress.configure(mode="determinate", maximum=len(all_xrefs), value=0)
        self.root.update()
        failed = []
        done = 0
        try:
            for idx, xref in enumerate(all_xrefs, 1):
                self.progress["value"] = idx - 1
                self.log(t("[{}/{}] {}").format(idx, len(all_xrefs), "xref" + str(xref)))
                self.root.update()
                try:
                    ext_img = doc.extract_image(xref)
                    data = ext_img["image"]
                    img = cv2.imdecode(np.frombuffer(data, np.uint8), cv2.IMREAD_COLOR)
                    if img is None:
                        self.log(t("  skip (cannot decode)"))
                        failed.append("xref" + str(xref))
                        continue
                    h, w = img.shape[:2]
                    abs_rects = [(rx1 * w, ry1 * h, rx2 * w, ry2 * h)
                                 for (rx1, ry1, rx2, ry2) in rel_rects]
                    extra = self._scale_mask_to(extra_ref, h0, w0, h, w)
                    mask = self._build_combined_mask(h, w, rects=abs_rects,
                                                     extra_mask=extra)
                    if mask.max() == 0:
                        self.log(t("  skip (empty mask)"))
                        failed.append("xref" + str(xref))
                        continue
                    result = self._run_iopaint(img, mask, BATCH_TEMP_ROOT)
                    ok, buf = cv2.imencode(".png", result)
                    if not ok:
                        self.log(t("  encode failed"))
                        failed.append("xref" + str(xref))
                        continue
                    # Replace image bytes in PDF.
                    replacement = fitz.Pixmap(buf.tobytes())
                    doc.update_stream(xref, buf.tobytes())
                    done += 1
                    self.log(t("  ok"))
                except Exception as e:
                    self.log(t("  FAILED: ") + str(e))
                    failed.append("xref" + str(xref))
            self.progress["value"] = len(all_xrefs)
            if done == 0:
                messagebox.showerror(t("Batch Remove"), t("No images were processed."))
                return
            doc.save(self.source_path, incremental=False, garbage=4, deflate=True,
                     clean=True)
            msg = t("Batch complete.\nProcessed: {}\nFailed: {}\nBackup: {}").format(
                done, len(failed), backup_path)
            self.log(msg)
            messagebox.showinfo(t("Batch Remove"), msg)
        finally:
            self._stop_progress()
            self.process_btn.config(state="normal")
            self.batch_btn.config(state="normal")
            self.root.config(cursor="")

    # -------------------- folder batch
    def batch_folder(self):
        if not self.rects and (self.extra_mask is None or self.extra_mask.max() == 0):
            messagebox.showwarning(
                t("Warning"), t("Draw at least one region on the current image first"))
            return
        if self.original_image is None:
            messagebox.showwarning(t("Warning"), t("Please open an image first"))
            return
        folder = filedialog.askdirectory(title=t("Select folder"))
        if not folder:
            return
        folder = Path(folder)
        files = [p for p in folder.iterdir()
                 if p.is_file() and p.suffix.lower() in IMAGE_EXTS]
        if not files:
            messagebox.showinfo(t("Info"), t("No images found"))
            return
        if not messagebox.askyesno(
            t("Batch Folder"),
            t("Process {} image(s) in {}?").format(len(files), str(folder))):
            return
        out_dir = folder / "_cleaned"
        out_dir.mkdir(exist_ok=True)
        h0, w0 = self.original_image.shape[:2]
        rel_rects = [(x1 / w0, y1 / h0, x2 / w0, y2 / h0)
                     for (x1, y1, x2, y2) in self.rects]
        extra_ref = self.extra_mask.copy() if self.extra_mask is not None else None
        self.progress.configure(mode="determinate", maximum=len(files), value=0)
        self.process_btn.config(state="disabled")
        self.batch_btn.config(state="disabled")
        self.root.config(cursor="watch")
        done = 0
        try:
            for i, fp in enumerate(files, 1):
                self.progress["value"] = i - 1
                self.log(t("[{}/{}] {}").format(i, len(files), fp.name))
                self.root.update()
                try:
                    img = cv2.imdecode(np.fromfile(str(fp), dtype=np.uint8),
                                       cv2.IMREAD_COLOR)
                    if img is None:
                        self.log(t("  skip (cannot decode)"))
                        continue
                    h, w = img.shape[:2]
                    abs_rects = [(rx1 * w, ry1 * h, rx2 * w, ry2 * h)
                                 for (rx1, ry1, rx2, ry2) in rel_rects]
                    extra = self._scale_mask_to(extra_ref, h0, w0, h, w)
                    mask = self._build_combined_mask(h, w, rects=abs_rects,
                                                     extra_mask=extra)
                    if mask.max() == 0:
                        self.log(t("  skip (empty mask)"))
                        continue
                    res = self._run_iopaint(img, mask, BATCH_TEMP_ROOT)
                    out_fp = out_dir / fp.name
                    ok, buf = cv2.imencode(fp.suffix.lower() or ".png", res)
                    if ok:
                        buf.tofile(str(out_fp))
                        done += 1
                        self.log(t("  ok"))
                except Exception as e:
                    self.log(t("  FAILED: ") + str(e))
            self.progress["value"] = len(files)
            messagebox.showinfo(
                t("Batch Folder"),
                t("Processed {} image(s). Output folder: {}").format(done, str(out_dir)))
        finally:
            self._stop_progress()
            self.process_btn.config(state="normal")
            self.batch_btn.config(state="normal")
            self.root.config(cursor="")

    # -------------------- auto-detect + OCR
    def auto_detect(self):
        if self.cv_image is None:
            messagebox.showwarning(t("Warning"), t("Please open an image first"))
            return
        cands = auto_detect_regions(self.cv_image)
        if not cands:
            self.log(t("No watermarks detected"))
            messagebox.showinfo(t("Info"), t("No watermarks detected"))
            return
        self.rects.extend(cands)
        self._redraw_overlay()
        self._update_coord_text()
        self.log(t("Detected {} candidate region(s)").format(len(cands)))

    def ocr_detect(self):
        if self.cv_image is None:
            messagebox.showwarning(t("Warning"), t("Please open an image first"))
            return
        if not HAVE_OCR:
            messagebox.showwarning(t("Warning"), t(
                "pytesseract / Tesseract not installed. Run: pip install pytesseract (and install Tesseract)"))
            return
        try:
            rgb = cv2.cvtColor(self.cv_image, cv2.COLOR_BGR2RGB)
            data = pytesseract.image_to_data(rgb, output_type=pytesseract.Output.DICT)
        except Exception as e:
            messagebox.showerror(t("Error"), t("Failed: ") + str(e))
            return
        added = 0
        for i in range(len(data.get("text", []))):
            txt = (data["text"][i] or "").strip()
            try:
                conf = int(data["conf"][i])
            except Exception:
                conf = 0
            if txt and conf >= 50:
                x = data["left"][i]
                y = data["top"][i]
                w = data["width"][i]
                h = data["height"][i]
                if w > 4 and h > 4:
                    self.rects.append((x, y, x + w, y + h))
                    added += 1
        if added == 0:
            messagebox.showinfo(t("Info"), t("No watermarks detected"))
            return
        self._redraw_overlay()
        self._update_coord_text()
        self.log(t("Detected {} candidate region(s)").format(added))

    # -------------------- save
    def save_to_source(self):
        if self.cv_image is None:
            messagebox.showerror(t("Error"), t("No image to save"))
            return
        if self.source_kind == "image" or self.source_kind is None:
            messagebox.showerror(
                t("Error"),
                t("This image was not opened from a document. Use Export PNG instead."))
            return
        try:
            if self.source_kind == "pdf":
                self._save_to_pdf()
            else:
                self._save_to_ooxml()
        except Exception as e:
            self.log(t("Failed to save: ") + str(e))
            messagebox.showerror(t("Error"), t("Failed to save: ") + str(e))

    def _save_to_ooxml(self):
        partname = self.source_extra
        backup_path = str(self.source_path) + ".backup"
        shutil.copy(self.source_path, backup_path)
        ext = Path(partname).suffix.lower() or ".png"
        ok, buffer = cv2.imencode(ext, self.cv_image)
        if not ok:
            raise RuntimeError("cv2.imencode failed")
        replace_ooxml_images(self.source_path, {partname: buffer.tobytes()}, EXTRACT_ROOT)
        if self.source_kind == "word":
            try:
                self.word_doc = Document(self.source_path)
            except Exception:
                pass
        self.log(t("Saved to {} (backup: {})").format(
            self.source_kind.upper(), backup_path))
        messagebox.showinfo(t("Success"), t("Saved to:\n") + self.source_path)

    def _save_to_pdf(self):
        if not HAVE_PDF or self.pdf_doc is None:
            messagebox.showwarning(
                t("Warning"), t("PyMuPDF not installed. Run: pip install pymupdf"))
            return
        xref = self.source_extra
        backup_path = str(self.source_path) + ".backup"
        self.pdf_doc.save(backup_path)
        ok, buf = cv2.imencode(".png", self.cv_image)
        if not ok:
            raise RuntimeError("cv2.imencode failed")
        try:
            self.pdf_doc.update_stream(xref, buf.tobytes())
        except Exception as e:
            raise RuntimeError("PDF image replace failed: " + str(e))
        self.pdf_doc.save(self.source_path, incremental=False, garbage=4,
                          deflate=True, clean=True)
        self.log(t("Saved to {} (backup: {})").format("PDF", backup_path))
        messagebox.showinfo(t("Success"), t("Saved to:\n") + self.source_path)

    def save_image(self):
        if self.cv_image is None:
            messagebox.showerror(t("Error"), t("No image to save"))
            return
        initial = CONFIG.get("last_dir") or None
        filepath = filedialog.asksaveasfilename(
            initialdir=initial,
            defaultextension=".png",
            filetypes=[("PNG files", "*.png"), ("JPEG files", "*.jpg"),
                       ("All files", "*.*")],
        )
        if not filepath:
            return
        ext = Path(filepath).suffix.lower() or ".png"
        ok, buf = cv2.imencode(ext, self.cv_image)
        if ok:
            buf.tofile(filepath)
            self.log(t("Exported: ") + Path(filepath).name)
            CONFIG["last_dir"] = str(Path(filepath).parent)
            save_config(CONFIG)
            messagebox.showinfo(t("Success"), t("Saved to:\n") + filepath)
        else:
            messagebox.showerror(t("Error"), t("Failed to encode image"))

    # -------------------- templates
    def save_template(self):
        if self.original_image is None or (not self.rects and (self.extra_mask is None or self.extra_mask.max() == 0)):
            messagebox.showwarning(t("Warning"), t("Draw at least one region first"))
            return
        filepath = filedialog.asksaveasfilename(
            defaultextension=TEMPLATE_EXT,
            filetypes=[("Watermark template", "*" + TEMPLATE_EXT),
                       ("All files", "*.*")],
            title=t("Save template"),
        )
        if not filepath:
            return
        h, w = self.original_image.shape[:2]
        rel_rects = [(x1 / w, y1 / h, x2 / w, y2 / h)
                     for (x1, y1, x2, y2) in self.rects]
        data = {
            "version": 1,
            "rel_rects": rel_rects,
            "model": self.model.get(),
            "feather": int(self.feather.get()),
            "color_tol": int(self.color_tol.get()),
        }
        if self.extra_mask is not None and self.extra_mask.max() > 0:
            # Encode mask as base64 PNG for portability.
            import base64
            ok, buf = cv2.imencode(".png", self.extra_mask)
            if ok:
                data["extra_mask_png"] = base64.b64encode(buf.tobytes()).decode("ascii")
                data["extra_mask_shape"] = list(self.extra_mask.shape)
        try:
            Path(filepath).write_text(
                json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
            self.log(t("Template saved") + ": " + filepath)
            messagebox.showinfo(t("Success"), t("Template saved"))
        except Exception as e:
            messagebox.showerror(t("Error"), t("Failed: ") + str(e))

    def load_template(self):
        filepath = filedialog.askopenfilename(
            filetypes=[("Watermark template", "*" + TEMPLATE_EXT),
                       ("All files", "*.*")],
            title=t("Load template"),
        )
        if filepath:
            self._apply_template_file(filepath)

    def _apply_template_file(self, filepath):
        if self.cv_image is None:
            messagebox.showwarning(t("Warning"), t("Please open an image first"))
            return
        try:
            data = json.loads(Path(filepath).read_text(encoding="utf-8"))
        except Exception as e:
            messagebox.showerror(t("Error"), t("Cannot load template: ") + str(e))
            return
        h, w = self.cv_image.shape[:2]
        self.rects = [(rx1 * w, ry1 * h, rx2 * w, ry2 * h)
                      for (rx1, ry1, rx2, ry2) in data.get("rel_rects", [])]
        if "extra_mask_png" in data:
            import base64
            raw = base64.b64decode(data["extra_mask_png"])
            nparr = np.frombuffer(raw, np.uint8)
            m = cv2.imdecode(nparr, cv2.IMREAD_GRAYSCALE)
            if m is not None:
                if m.shape != (h, w):
                    m = cv2.resize(m, (w, h), interpolation=cv2.INTER_NEAREST)
                self.extra_mask = m
        self.model.set(data.get("model", "lama"))
        self.feather.set(int(data.get("feather", 0)))
        self.color_tol.set(int(data.get("color_tol", 24)))
        self.display_image()
        self._update_coord_text()
        self.log(t("Template loaded"))

    # -------------------- compare
    def show_compare(self):
        if self.original_image is None or self.cv_image is None:
            messagebox.showwarning(t("Warning"), t("Please open an image first"))
            return
        top = tk.Toplevel(self.root)
        top.title(t("Compare"))
        top.geometry("1100x640")

        def _to_photo(img, maxw, maxh):
            h, w = img.shape[:2]
            s = min(maxw / w, maxh / h, 1.0)
            nw, nh = max(1, int(w * s)), max(1, int(h * s))
            rgb = cv2.cvtColor(cv2.resize(img, (nw, nh)), cv2.COLOR_BGR2RGB)
            return ImageTk.PhotoImage(Image.fromarray(rgb))

        left_frame = tk.Frame(top)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        tk.Label(left_frame, text=t("Before"),
                 font=("Arial", 11, "bold")).pack()
        l = tk.Label(left_frame)
        l.pack(fill=tk.BOTH, expand=True)
        right_frame = tk.Frame(top)
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        tk.Label(right_frame, text=t("After"),
                 font=("Arial", 11, "bold")).pack()
        r = tk.Label(right_frame)
        r.pack(fill=tk.BOTH, expand=True)

        top.update()
        p1 = _to_photo(self.original_image, 540, 580)
        p2 = _to_photo(self.cv_image, 540, 580)
        l.configure(image=p1)
        r.configure(image=p2)
        l.image = p1
        r.image = p2

    # -------------------- platform templates
    def apply_platform_template(self, name):
        if self.cv_image is None:
            messagebox.showwarning(t("Warning"), t("Please open an image first"))
            return
        tpl = PLATFORM_TEMPLATES.get(name)
        if not tpl:
            return
        h, w = self.cv_image.shape[:2]
        for rx1, ry1, rx2, ry2 in tpl:
            self.rects.append((rx1 * w, ry1 * h, rx2 * w, ry2 * h))
        self._redraw_overlay()
        self._update_coord_text()
        self.log(t("Template applied") + ": " + name)

    # -------------------- loupe
    def toggle_loupe(self):
        if self.loupe_win is not None and self.loupe_win.winfo_exists():
            self.loupe_win.destroy()
            self.loupe_win = None
            self.loupe_canvas = None
            return
        if self.cv_image is None:
            messagebox.showwarning(t("Warning"), t("Please open an image first"))
            return
        self.loupe_win = tk.Toplevel(self.root)
        self.loupe_win.title(t("Loupe"))
        self.loupe_win.geometry("360x360+{}+{}".format(
            self.root.winfo_x() + self.root.winfo_width() - 380,
            self.root.winfo_y() + 40,
        ))
        self.loupe_win.resizable(False, False)
        self.loupe_win.attributes("-topmost", True)
        self.loupe_canvas = tk.Canvas(self.loupe_win, bg="black",
                                      width=340, height=340)
        self.loupe_canvas.pack(padx=5, pady=5)

    def _update_loupe(self, ix, iy):
        if self.loupe_win is None or not self.loupe_win.winfo_exists():
            self.loupe_win = None
            self.loupe_canvas = None
            return
        if self.cv_image is None:
            return
        h, w = self.cv_image.shape[:2]
        crop = 60
        x0 = max(0, int(ix) - crop)
        y0 = max(0, int(iy) - crop)
        x1 = min(w, int(ix) + crop)
        y1 = min(h, int(iy) + crop)
        if x1 <= x0 or y1 <= y0:
            return
        patch = self.cv_image[y0:y1, x0:x1]
        if patch.size == 0:
            return
        # 3x zoom
        zoomed = cv2.resize(patch, (340, 340), interpolation=cv2.INTER_NEAREST)
        rgb = cv2.cvtColor(zoomed, cv2.COLOR_BGR2RGB)
        # Crosshair in the middle
        cv2.line(rgb, (170, 0), (170, 340), (255, 255, 0), 1)
        cv2.line(rgb, (0, 170), (340, 170), (255, 255, 0), 1)
        self.loupe_photo = ImageTk.PhotoImage(Image.fromarray(rgb))
        self.loupe_canvas.delete("all")
        self.loupe_canvas.create_image(170, 170, image=self.loupe_photo)

    # -------------------- remove preview + deep clean
    def _preview_remove_result(self, new_image):
        """Show a modal side-by-side; returns True if user hits Apply."""
        top = tk.Toplevel(self.root)
        top.title(t("Preview result"))
        top.geometry("1200x720")
        top.transient(self.root)
        top.grab_set()

        def _photo_of(img, maxw, maxh):
            h, w = img.shape[:2]
            s = min(maxw / w, maxh / h, 1.0)
            nw, nh = max(1, int(w * s)), max(1, int(h * s))
            rgb = cv2.cvtColor(cv2.resize(img, (nw, nh)), cv2.COLOR_BGR2RGB)
            return ImageTk.PhotoImage(Image.fromarray(rgb))

        body = tk.Frame(top)
        body.pack(fill=tk.BOTH, expand=True, padx=8, pady=8)
        lf = tk.Frame(body)
        lf.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=4)
        rf = tk.Frame(body)
        rf.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=4)
        tk.Label(lf, text=t("Before"), font=("Arial", 11, "bold")).pack()
        ll = tk.Label(lf)
        ll.pack(fill=tk.BOTH, expand=True)
        tk.Label(rf, text=t("After"), font=("Arial", 11, "bold")).pack()
        rl = tk.Label(rf)
        rl.pack(fill=tk.BOTH, expand=True)
        top.update()
        p1 = _photo_of(self.cv_image, 580, 620)
        p2 = _photo_of(new_image, 580, 620)
        ll.configure(image=p1)
        rl.configure(image=p2)
        ll.image = p1
        rl.image = p2

        decision = {"apply": False}

        def _apply():
            decision["apply"] = True
            top.destroy()

        def _cancel():
            decision["apply"] = False
            top.destroy()

        bar = tk.Frame(top)
        bar.pack(fill=tk.X, pady=6)
        tk.Button(bar, text=t("Apply"), command=_apply, bg="lightgreen",
                  font=("Arial", 11, "bold"), width=14).pack(side=tk.RIGHT, padx=8)
        tk.Button(bar, text=t("Cancel"), command=_cancel, bg="lightgray",
                  width=14).pack(side=tk.RIGHT, padx=4)
        top.protocol("WM_DELETE_WINDOW", _cancel)
        top.wait_window()
        return decision["apply"]

    def deep_clean(self, max_passes=3):
        if self.cv_image is None:
            messagebox.showerror(t("Error"), t("Please open an image first"))
            return
        # Either use existing selection OR seed from auto-detect.
        if not self.rects and (self.extra_mask is None or self.extra_mask.max() == 0):
            cands = auto_detect_regions(self.cv_image)
            if not cands:
                messagebox.showinfo(t("Info"), t("No watermarks detected"))
                return
            self.rects.extend(cands)
            self._redraw_overlay()

        for pass_idx in range(1, max_passes + 1):
            self.log(t("Deep clean pass {}/{}").format(pass_idx, max_passes))
            if not self.rects and (self.extra_mask is None or self.extra_mask.max() == 0):
                self.log(t("Nothing detected after pass {}. Stopping.").format(pass_idx - 1))
                break
            if not self._remove_core(show_preview=False):
                break
            new_cands = auto_detect_regions(self.cv_image)
            if not new_cands:
                self.log(t("Nothing detected after pass {}. Stopping.").format(pass_idx))
                break
            self.rects.extend(new_cands)
        self.display_image()
        self._update_coord_text()

    # -------------------- video
    def open_video(self):
        initial = CONFIG.get("last_dir") or None
        fp = filedialog.askopenfilename(
            initialdir=initial,
            filetypes=[("Video", " ".join("*" + e for e in VIDEO_EXTS)),
                       ("All files", "*.*")],
        )
        if not fp:
            return
        cap = cv2.VideoCapture(fp)
        if not cap.isOpened():
            messagebox.showerror(t("Error"), t("Cannot open video"))
            return
        ok, frame = cap.read()
        cap.release()
        if not ok or frame is None:
            messagebox.showerror(t("Error"), t("Cannot open video"))
            return
        self.cv_image = frame
        self.original_image = frame.copy()
        self.source_kind = "image"
        self.source_path = fp
        self.video_path = fp
        self._file_desc = ("image", Path(fp).name + "  (video frame 0)")
        self._reset_selection_state()
        self._reset_undo()
        self._refresh_file_text()
        self.log(t("Opened: ") + Path(fp).name)
        self.log(t("Size: {}x{}").format(frame.shape[1], frame.shape[0]))
        self.reset_view()
        self._remember(fp)

    def process_video(self):
        if not self.video_path:
            messagebox.showwarning(t("Warning"), t("Open a video first via Video menu"))
            return
        if not self.rects and (self.extra_mask is None or self.extra_mask.max() == 0):
            messagebox.showwarning(t("Warning"),
                                   t("Draw at least one region first"))
            return
        out_path = filedialog.asksaveasfilename(
            title=t("Export video to:"),
            defaultextension=".mp4",
            filetypes=[("MP4", "*.mp4"), ("AVI", "*.avi"), ("All", "*.*")],
        )
        if not out_path:
            return

        cap = cv2.VideoCapture(self.video_path)
        fps = cap.get(cv2.CAP_PROP_FPS) or 30.0
        total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
        w = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
        h = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
        fourcc = cv2.VideoWriter_fourcc(*"mp4v")
        writer = cv2.VideoWriter(out_path, fourcc, fps, (w, h))
        if not writer.isOpened():
            cap.release()
            messagebox.showerror(t("Error"), t("Cannot open video"))
            return

        # Build mask once (scales to each frame size in principle, but videos
        # have constant dims so we reuse it).
        h0, w0 = self.original_image.shape[:2]
        rel_rects = [(x1 / w0, y1 / h0, x2 / w0, y2 / h0)
                     for (x1, y1, x2, y2) in self.rects]
        abs_rects = [(rx1 * w, ry1 * h, rx2 * w, ry2 * h)
                     for rx1, ry1, rx2, ry2 in rel_rects]
        extra = self._scale_mask_to(self.extra_mask, h0, w0, h, w) \
            if self.extra_mask is not None else None
        mask = self._build_combined_mask(h, w, rects=abs_rects, extra_mask=extra)
        if mask.max() == 0:
            cap.release()
            writer.release()
            messagebox.showerror(t("Error"), t("Selection is empty"))
            return

        self.progress.configure(mode="determinate",
                                maximum=max(1, total), value=0)
        self.process_btn.config(state="disabled")
        self.batch_btn.config(state="disabled")
        self.root.config(cursor="watch")
        done = 0
        try:
            # Fast path: OpenCV Telea inpainting. Much faster than LaMa and
            # good enough for typical static watermarks on video. Running
            # LaMa per frame would take hours.
            while True:
                ok, frame = cap.read()
                if not ok:
                    break
                cleaned = cv2.inpaint(frame, mask, 3, cv2.INPAINT_TELEA)
                writer.write(cleaned)
                done += 1
                if done % 10 == 0:
                    self.progress["value"] = done
                    self.log(t("Processing frame {}/{}").format(done, total))
                    self.root.update()
            self.progress["value"] = done
            self.log(t("Video done. Output: {}").format(out_path))
            messagebox.showinfo(t("Success"),
                                t("Video done. Output: {}").format(out_path))
        finally:
            cap.release()
            writer.release()
            self._stop_progress()
            self.process_btn.config(state="normal")
            self.batch_btn.config(state="normal")
            self.root.config(cursor="")

    # -------------------- closing
    def _on_close(self):
        try:
            CONFIG["geometry"] = self.root.winfo_geometry()
            CONFIG["model"] = self.model.get()
            CONFIG["feather"] = int(self.feather.get())
            CONFIG["brush_size"] = int(self.brush_size.get())
            CONFIG["color_tol"] = int(self.color_tol.get())
            save_config(CONFIG)
        except Exception:
            pass
        self.root.destroy()


# ----------------------------------------------------------- CLI
def run_cli(argv):
    ap = argparse.ArgumentParser(description="Watermark remover (headless mode).")
    ap.add_argument("--image", required=True)
    ap.add_argument("--mask", required=True, help="grayscale mask PNG (255 = remove)")
    ap.add_argument("--output", required=True)
    ap.add_argument("--model", default="lama")
    args = ap.parse_args(argv)
    img = cv2.imdecode(np.fromfile(args.image, dtype=np.uint8), cv2.IMREAD_COLOR)
    mask = cv2.imdecode(np.fromfile(args.mask, dtype=np.uint8), cv2.IMREAD_GRAYSCALE)
    if img is None or mask is None:
        print("Cannot read input", file=sys.stderr)
        return 1
    if mask.shape != img.shape[:2]:
        mask = cv2.resize(mask, (img.shape[1], img.shape[0]),
                          interpolation=cv2.INTER_NEAREST)
    tmp = TEMP_ROOT / ("cli-" + str(int(time.time())))
    tmp.mkdir(exist_ok=True, parents=True)
    ip = tmp / "input.png"
    mp = tmp / "mask.png"
    out = tmp / "output"
    out.mkdir(exist_ok=True)
    cv2.imwrite(str(ip), img)
    cv2.imwrite(str(mp), mask)
    exe = sys.executable
    if getattr(sys, "frozen", False):
        cmd = [exe, "--iopaint-worker", "run", "--model", args.model,
               "--image", str(ip), "--mask", str(mp), "--output", str(out)]
    else:
        cmd = [exe, "-m", "iopaint", "run", "--model", args.model,
               "--image", str(ip), "--mask", str(mp), "--output", str(out)]
    subprocess.run(cmd, check=True)
    result = cv2.imread(str(out / "input.png"))
    if result is None:
        print("iopaint produced no output", file=sys.stderr)
        return 2
    ext = Path(args.output).suffix.lower() or ".png"
    ok, buf = cv2.imencode(ext, result)
    if not ok:
        print("Failed to encode output", file=sys.stderr)
        return 3
    buf.tofile(args.output)
    return 0


# ----------------------------------------------------------- main
if __name__ == "__main__":
    import multiprocessing
    multiprocessing.freeze_support()

    if len(sys.argv) > 1 and sys.argv[1] == "--iopaint-worker":
        from iopaint import entry_point
        sys.argv = ["iopaint"] + sys.argv[2:]
        entry_point()
        sys.exit(0)

    if len(sys.argv) > 1 and sys.argv[1] == "--cli":
        sys.exit(run_cli(sys.argv[2:]))

    if HAVE_DND:
        root = TkinterDnD.Tk()
    else:
        root = tk.Tk()
    app = WatermarkRemover(root)
    root.mainloop()
